"""
Generate 9 CVLPOS contract templates as .docx files.
Uses python-docx to create Word documents with {{ variable }} placeholders
that docxtpl will render later.
"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_DIR = "/Users/yasudaosamu/Desktop/codes/auction/templates/contracts"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_run_font(run, size_pt=10.5, bold=False, font_name="游明朝"):
    """Configure a run with Japanese-friendly font settings."""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = font_name
    # Set East-Asian font via XML so Word picks it up correctly
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(text)
    set_run_font(run, size_pt=16, bold=True)


def add_heading_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size_pt=12, bold=True)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size_pt=10.5)
    return p


def add_blank_line(doc):
    doc.add_paragraph()


def add_signature_block(doc):
    add_blank_line(doc)
    add_blank_line(doc)
    add_body(doc, "甲 署名: ____________________  印")
    add_blank_line(doc)
    add_body(doc, "乙 署名: ____________________  印")


def set_default_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ---------------------------------------------------------------------------
# 1. 匿名組合契約書 (TK Agreement)
# ---------------------------------------------------------------------------

def create_tk_agreement():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "匿名組合契約書")

    # Parties
    add_heading_text(doc, "当事者")
    add_body(doc, "甲（営業者）: {{ party_a_name }}")
    add_body(doc, "住所: {{ party_a_address }}")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（匿名組合員）: {{ party_b_name }}")
    add_body(doc, "住所: {{ party_b_address }}")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")

    # Article 1
    add_heading_text(doc, "第1条（目的）")
    add_body(doc,
        "甲は、乙から出資を受けた金員をもって、商用車両のリースバック事業"
        "（以下「本件事業」という）を営み、その営業から生ずる利益を乙に分配する。")

    # Article 2
    add_heading_text(doc, "第2条（出資）")
    add_body(doc, "乙は、本契約に基づき、金{{ purchase_price }}円を甲に出資する。")
    add_body(doc,
        "出資金は、本契約締結後速やかに甲の指定する口座に振り込む方法により払い込む。")

    # Article 3
    add_heading_text(doc, "第3条（利益分配）")
    add_body(doc,
        "甲は、本件事業から生じる利益を、目標利回り{{ target_yield_rate }}に基づき乙に分配する。")
    add_body(doc,
        "分配は月額{{ monthly_lease_fee }}円のリース料収入を原資とする。")
    add_body(doc,
        "分配金は毎月末日に計算し、翌月{{ payment_day }}日までに乙の指定口座に振り込む。")

    # Article 4
    add_heading_text(doc, "第4条（契約期間）")
    add_body(doc, "本契約の期間は、{{ lease_term_months }}ヶ月とする。")
    add_body(doc,
        "開始日: {{ lease_start_date }}  終了日: {{ lease_end_date }}")

    # Article 5
    add_heading_text(doc, "第5条（対象車両）")
    add_body(doc, "メーカー: {{ vehicle_maker }}")
    add_body(doc, "車種: {{ vehicle_model }}")
    add_body(doc, "年式: {{ vehicle_year }}年")
    add_body(doc, "走行距離: {{ vehicle_mileage }}")
    add_body(doc, "車台番号: {{ vehicle_chassis_number }}")
    add_body(doc, "登録番号: {{ vehicle_registration_number }}")

    # Article 6
    add_heading_text(doc, "第6条（業務執行）")
    add_body(doc,
        "甲は、善良なる管理者の注意をもって本件事業の業務を執行する。"
        "乙は、甲の業務執行に対し指示を行い、又はこれに参加することはできない。")

    # Article 7
    add_heading_text(doc, "第7条（解約・終了）")
    add_body(doc,
        "本契約は、契約期間の満了をもって終了する。"
        "やむを得ない事由がある場合は、3ヶ月前の書面通知により中途解約できるものとする。")

    # Article 8
    add_heading_text(doc, "第8条（秘密保持）")
    add_body(doc,
        "甲及び乙は、本契約に関して知り得た相手方の秘密情報を、"
        "相手方の事前の書面による承諾なく第三者に開示又は漏洩してはならない。")

    # Article 9
    add_heading_text(doc, "第9条（管轄裁判所）")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "tk_agreement.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# 2. 車両売買契約書 (Sales Agreement)
# ---------------------------------------------------------------------------

def create_sales_agreement():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "車両売買契約書")

    add_heading_text(doc, "当事者")
    add_body(doc, "甲（売主）: {{ party_a_name }} ({{ party_a_address }})")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（買主）: {{ party_b_name }} ({{ party_b_address }})")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")

    # Article 1
    add_heading_text(doc, "第1条（売買の合意）")
    add_body(doc,
        "甲は、乙に対し、末尾記載の車両を金{{ purchase_price }}円（消費税別）にて売却し、"
        "乙はこれを買い受ける。")

    # Article 2
    add_heading_text(doc, "第2条（対象車両）")
    add_body(doc, "メーカー: {{ vehicle_maker }}")
    add_body(doc, "車種: {{ vehicle_model }} / 年式: {{ vehicle_year }}年")
    add_body(doc, "走行距離: {{ vehicle_mileage }}")
    add_body(doc, "架装: {{ vehicle_body_type }}")
    add_body(doc, "車台番号: {{ vehicle_chassis_number }}")
    add_body(doc, "登録番号: {{ vehicle_registration_number }}")

    # Article 3
    add_heading_text(doc, "第3条（代金支払い）")
    add_body(doc,
        "乙は、本契約締結後速やかに売買代金を甲の指定する口座に振り込むものとする。")
    add_body(doc,
        "振込先: {{ payment_bank_name }} {{ payment_branch_name }} "
        "{{ payment_account_type }} {{ payment_account_number }}")

    # Article 4
    add_heading_text(doc, "第4条（名義変更・引渡し）")
    add_body(doc,
        "甲は、売買代金の受領を確認後、速やかに車両の名義変更手続を行い、"
        "乙に対して車両を引き渡すものとする。")
    add_body(doc,
        "名義変更に要する費用は、乙の負担とする。")

    # Article 5
    add_heading_text(doc, "第5条（瑕疵担保）")
    add_body(doc,
        "甲は、引渡し後{{ warranty_months }}ヶ月間、対象車両の隠れた瑕疵について担保責任を負う。"
        "ただし、乙の故意又は過失による損傷についてはこの限りでない。")

    # Article 6
    add_heading_text(doc, "第6条（危険負担）")
    add_body(doc,
        "車両の引渡し前に生じた滅失・毀損の危険は甲が負担し、"
        "引渡し後の危険は乙が負担する。")

    # Article 7
    add_heading_text(doc, "第7条（契約解除）")
    add_body(doc,
        "甲又は乙が本契約に違反し、相当の期間を定めて催告したにもかかわらず"
        "是正されない場合、相手方は本契約を解除することができる。")

    # Article 8
    add_heading_text(doc, "第8条（管轄裁判所）")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "sales_agreement.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# 3. マスターリース契約書 (Master Lease)
# ---------------------------------------------------------------------------

def create_master_lease():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "マスターリース契約書")

    add_heading_text(doc, "当事者")
    add_body(doc, "甲（貸人）: {{ party_a_name }}")
    add_body(doc, "住所: {{ party_a_address }}")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（借人）: {{ party_b_name }}")
    add_body(doc, "住所: {{ party_b_address }}")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")

    # Article 1
    add_heading_text(doc, "第1条（リースの合意）")
    add_body(doc,
        "甲は、乙に対し末尾記載の車両をリースし、乙はこれを借り受ける。")

    # Article 2
    add_heading_text(doc, "第2条（リース料）")
    add_body(doc, "月額リース料: {{ monthly_lease_fee }}円（消費税別）")
    add_body(doc, "リース期間: {{ lease_term_months }}ヶ月")
    add_body(doc, "リース料総額: {{ total_lease_revenue }}円")
    add_body(doc,
        "乙は、毎月{{ payment_day }}日までに当月分のリース料を甲の指定口座に振り込むものとする。")

    # Article 3
    add_heading_text(doc, "第3条（対象車両）")
    add_body(doc, "メーカー: {{ vehicle_maker }}")
    add_body(doc, "車種: {{ vehicle_model }}")
    add_body(doc, "年式: {{ vehicle_year }}年")
    add_body(doc, "走行距離: {{ vehicle_mileage }}")
    add_body(doc, "車台番号: {{ vehicle_chassis_number }}")
    add_body(doc, "登録番号: {{ vehicle_registration_number }}")

    # Article 4
    add_heading_text(doc, "第4条（使用目的）")
    add_body(doc,
        "乙は、対象車両を営業用車両として使用するものとし、"
        "甲の事前の書面による承諾なく第三者に転貸してはならない。"
        "ただし、第三者へのサブリースについて甲が書面で承諾した場合はこの限りでない。")

    # Article 5
    add_heading_text(doc, "第5条（維持管理）")
    add_body(doc,
        "乙は、善良なる管理者の注意をもって対象車両を使用・保管し、"
        "通常の維持管理費用（点検・整備・保険等）を負担する。")

    # Article 6
    add_heading_text(doc, "第6条（保険）")
    add_body(doc,
        "乙は、リース期間中、対象車両について自動車保険（対人・対物無制限）に加入し、"
        "その証書の写しを甲に提出するものとする。")

    # Article 7
    add_heading_text(doc, "第7条（中途解約）")
    add_body(doc,
        "本契約は、リース期間中の中途解約はできないものとする。"
        "ただし、やむを得ない事由がある場合は、残リース料相当額を違約金として支払うことにより"
        "解約できるものとする。")

    # Article 8
    add_heading_text(doc, "第8条（契約終了時の処理）")
    add_body(doc,
        "リース期間満了時、乙は対象車両を原状に回復のうえ甲に返還するものとする。")

    # Article 9
    add_heading_text(doc, "第9条（管轄裁判所）")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "master_lease.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# 4. サブリース契約書 (Sublease Agreement)
# ---------------------------------------------------------------------------

def create_sublease_agreement():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "サブリース契約書")

    add_heading_text(doc, "当事者")
    add_body(doc, "甲（転貸人）: {{ party_a_name }}")
    add_body(doc, "住所: {{ party_a_address }}")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（転借人）: {{ party_b_name }}")
    add_body(doc, "住所: {{ party_b_address }}")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")

    # Article 1
    add_heading_text(doc, "第1条（サブリースの合意）")
    add_body(doc,
        "甲は、甲がマスターリース契約に基づきリースを受けている末尾記載の車両を、"
        "乙に対しサブリース（転貸）し、乙はこれを借り受ける。")

    # Article 2
    add_heading_text(doc, "第2条（サブリース料）")
    add_body(doc, "月額サブリース料: {{ sublease_fee }}円（消費税別）")
    add_body(doc, "サブリース期間: {{ lease_term_months }}ヶ月")
    add_body(doc,
        "乙は、毎月{{ payment_day }}日までに当月分のサブリース料を甲の指定口座に振り込むものとする。")

    # Article 3
    add_heading_text(doc, "第3条（対象車両）")
    add_body(doc, "メーカー: {{ vehicle_maker }}")
    add_body(doc, "車種: {{ vehicle_model }}")
    add_body(doc, "年式: {{ vehicle_year }}年")
    add_body(doc, "走行距離: {{ vehicle_mileage }}")
    add_body(doc, "車台番号: {{ vehicle_chassis_number }}")
    add_body(doc, "登録番号: {{ vehicle_registration_number }}")

    # Article 4
    add_heading_text(doc, "第4条（使用条件）")
    add_body(doc,
        "乙は、対象車両を事業用車両として使用するものとし、"
        "善良なる管理者の注意をもってこれを使用・保管する。")
    add_body(doc,
        "乙は、甲の事前の書面による承諾なく、対象車両を第三者に再転貸してはならない。")

    # Article 5
    add_heading_text(doc, "第5条（維持管理・保険）")
    add_body(doc,
        "乙は、サブリース期間中の維持管理費用（燃料費、点検・整備費用等）を負担する。")
    add_body(doc,
        "乙は、対象車両について自動車保険（対人・対物無制限）に加入し、"
        "その証書の写しを甲に提出するものとする。")

    # Article 6
    add_heading_text(doc, "第6条（事故・故障時の対応）")
    add_body(doc,
        "乙は、対象車両に事故又は故障が生じた場合、直ちに甲に報告し、"
        "甲の指示に従って対応するものとする。")

    # Article 7
    add_heading_text(doc, "第7条（中途解約）")
    add_body(doc,
        "本契約は、サブリース期間中の中途解約はできないものとする。"
        "ただし、甲乙協議のうえ合意した場合はこの限りでない。")

    # Article 8
    add_heading_text(doc, "第8条（契約終了時の処理）")
    add_body(doc,
        "サブリース期間満了時、乙は対象車両を原状に回復のうえ甲に返還するものとする。")

    # Article 9
    add_heading_text(doc, "第9条（マスターリースとの関係）")
    add_body(doc,
        "本サブリース契約は、甲と車両所有者との間のマスターリース契約に従属する。"
        "マスターリース契約が終了した場合、本契約も当然に終了するものとする。")

    # Article 10
    add_heading_text(doc, "第10条（管轄裁判所）")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "sublease_agreement.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# 5. 私募取扱業務契約書 (Private Placement Agreement)
# ---------------------------------------------------------------------------

def create_private_placement_agreement():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "私募取扱業務契約書")

    # 前文
    add_body(doc,
        "{{ party_a_name }}（以下「甲」という）と{{ party_b_name }}（以下「乙」という）は、"
        "甲が組成するファンドに係る私募の取扱業務に関し、以下のとおり契約を締結する。")

    # Parties
    add_heading_text(doc, "当事者")
    add_body(doc, "甲（発行者）: {{ party_a_name }}")
    add_body(doc, "住所: {{ party_a_address }}")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（取扱者）: {{ party_b_name }}")
    add_body(doc, "住所: {{ party_b_address }}")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")
    add_body(doc, "効力発生日: {{ effective_date }}")

    # Article 1
    add_heading_text(doc, "第1条（目的）")
    add_body(doc,
        "本契約は、甲が組成する{{ fund_name }}（以下「本ファンド」という）の持分の私募の取扱いに関し、"
        "甲が乙に対し当該業務を委託し、乙がこれを受託することを目的とする。")

    # Article 2
    add_heading_text(doc, "第2条（業務内容）")
    add_body(doc,
        "乙は、甲の指示に基づき、本ファンドの持分について適格機関投資家等への私募の取扱いを行う。")
    add_body(doc,
        "乙は、金融商品取引法その他関連法令を遵守し、適切な方法により私募の取扱いを行うものとする。")

    # Article 3
    add_heading_text(doc, "第3条（募集総額）")
    add_body(doc, "本ファンドの募集総額: 金{{ total_amount }}円")

    # Article 4
    add_heading_text(doc, "第4条（報酬）")
    add_body(doc,
        "甲は、乙に対し、私募取扱業務の対価として、取扱金額に対し{{ placement_fee_rate }}を乗じた金額"
        "（以下「取扱手数料」という）を支払う。")
    add_body(doc, "取扱手数料: 金{{ placement_fee_amount }}円")
    add_body(doc,
        "取扱手数料は、募集完了後30日以内に乙の指定する口座に振り込む方法により支払う。")

    # Article 5
    add_heading_text(doc, "第5条（期間）")
    add_body(doc,
        "本契約の有効期間は、効力発生日から1年間とする。"
        "ただし、期間満了の1ヶ月前までに甲乙いずれからも書面による終了の申出がないときは、"
        "同一条件で更に1年間更新されるものとし、以後も同様とする。")

    # Article 6
    add_heading_text(doc, "第6条（善管注意義務）")
    add_body(doc,
        "乙は、善良なる管理者の注意をもって本契約に基づく業務を遂行するものとする。")

    # Article 7
    add_heading_text(doc, "第7条（秘密保持）")
    add_body(doc,
        "甲及び乙は、本契約に関して知り得た相手方の秘密情報を、"
        "相手方の事前の書面による承諾なく第三者に開示又は漏洩してはならない。"
        "本条の義務は、本契約終了後もなお3年間存続する。")

    # Article 8
    add_heading_text(doc, "第8条（解除）")
    add_body(doc,
        "甲又は乙は、相手方が本契約の条項に違反し、相当の期間を定めて催告したにもかかわらず"
        "是正されない場合、本契約を解除することができる。")

    # Article 9
    add_heading_text(doc, "第9条（反社会的勢力の排除）")
    add_body(doc,
        "甲及び乙は、自ら又はその役員等が反社会的勢力に該当しないことを表明し、保証する。")

    # Article 10
    add_heading_text(doc, "第10条（準拠法及び管轄裁判所）")
    add_body(doc, "本契約は日本法に準拠し、日本法に従い解釈されるものとする。")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    # 後文
    add_blank_line(doc)
    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "private_placement_agreement.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# 6. 顧客紹介業務契約書 (Customer Referral Agreement)
# ---------------------------------------------------------------------------

def create_customer_referral_agreement():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "顧客紹介業務契約書")

    # 前文
    add_body(doc,
        "{{ party_a_name }}（以下「甲」という）と{{ party_b_name }}（以下「乙」という）は、"
        "甲が運営するファンドへの投資家紹介業務に関し、以下のとおり契約を締結する。")

    # Parties
    add_heading_text(doc, "当事者")
    add_body(doc, "甲（委託者）: {{ party_a_name }}")
    add_body(doc, "住所: {{ party_a_address }}")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（紹介者）: {{ party_b_name }}")
    add_body(doc, "住所: {{ party_b_address }}")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")
    add_body(doc, "効力発生日: {{ effective_date }}")

    # Article 1
    add_heading_text(doc, "第1条（目的）")
    add_body(doc,
        "本契約は、甲が運営する{{ fund_name }}（以下「本ファンド」という）への投資を検討する"
        "顧客の紹介業務に関し、甲が乙に対し当該業務を委託し、乙がこれを受託することを目的とする。")

    # Article 2
    add_heading_text(doc, "第2条（業務内容）")
    add_body(doc, "紹介業務の範囲: {{ referral_scope }}")
    add_body(doc,
        "乙は、本ファンドへの投資に関心を有する見込顧客を甲に紹介するものとする。")
    add_body(doc,
        "乙は、紹介にあたり、金融商品取引法上の勧誘行為に該当する行為を行ってはならない。")

    # Article 3
    add_heading_text(doc, "第3条（報酬）")
    add_body(doc,
        "甲は、乙が紹介した顧客が本ファンドへの出資を行った場合、"
        "当該出資金額に{{ referral_fee_rate }}を乗じた金額を紹介報酬として乙に支払う。")
    add_body(doc,
        "紹介報酬は、出資金の払込確認後30日以内に乙の指定する口座に振り込む方法により支払う。")

    # Article 4
    add_heading_text(doc, "第4条（期間）")
    add_body(doc,
        "本契約の有効期間は、効力発生日から1年間とする。"
        "ただし、期間満了の1ヶ月前までに甲乙いずれからも書面による終了の申出がないときは、"
        "同一条件で更に1年間更新されるものとし、以後も同様とする。")

    # Article 5
    add_heading_text(doc, "第5条（禁止事項）")
    add_body(doc,
        "乙は、以下の行為を行ってはならない。")
    add_body(doc, "（1）甲の名義を使用して契約を締結すること")
    add_body(doc, "（2）本ファンドの運用内容について断定的判断を提供すること")
    add_body(doc, "（3）虚偽又は誤解を招く情報を顧客に提供すること")

    # Article 6
    add_heading_text(doc, "第6条（秘密保持）")
    add_body(doc,
        "甲及び乙は、本契約に関して知り得た相手方の秘密情報を、"
        "相手方の事前の書面による承諾なく第三者に開示又は漏洩してはならない。"
        "本条の義務は、本契約終了後もなお3年間存続する。")

    # Article 7
    add_heading_text(doc, "第7条（解除）")
    add_body(doc,
        "甲又は乙は、相手方が本契約の条項に違反し、相当の期間を定めて催告したにもかかわらず"
        "是正されない場合、本契約を解除することができる。")

    # Article 8
    add_heading_text(doc, "第8条（損害賠償）")
    add_body(doc,
        "甲又は乙が本契約に違反し、相手方に損害を与えた場合、"
        "その損害を賠償する責任を負うものとする。")

    # Article 9
    add_heading_text(doc, "第9条（準拠法及び管轄裁判所）")
    add_body(doc, "本契約は日本法に準拠し、日本法に従い解釈されるものとする。")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    # 後文
    add_blank_line(doc)
    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "customer_referral_agreement.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# 7. アセットマネジメント契約書 (Asset Management Agreement)
# ---------------------------------------------------------------------------

def create_asset_management_agreement():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "アセットマネジメント契約書")

    # 前文
    add_body(doc,
        "{{ party_a_name }}（以下「甲」という）と{{ party_b_name }}（以下「乙」という）は、"
        "甲が保有する資産の運用管理業務に関し、以下のとおり契約を締結する。")

    # Parties
    add_heading_text(doc, "当事者")
    add_body(doc, "甲（委託者）: {{ party_a_name }}")
    add_body(doc, "住所: {{ party_a_address }}")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（受託者）: {{ party_b_name }}")
    add_body(doc, "住所: {{ party_b_address }}")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")
    add_body(doc, "効力発生日: {{ effective_date }}")

    # Article 1
    add_heading_text(doc, "第1条（目的）")
    add_body(doc,
        "本契約は、{{ fund_name }}（以下「本ファンド」という）に係る資産の運用管理業務を、"
        "甲が乙に委託し、乙がこれを受託することを目的とする。")

    # Article 2
    add_heading_text(doc, "第2条（業務内容）")
    add_body(doc,
        "乙は、甲の指示及び本契約の定めに従い、以下の業務を行う。")
    add_body(doc, "（1）運用資産の取得、処分及び管理に関する助言・実行")
    add_body(doc, "（2）運用資産の価値評価及びモニタリング")
    add_body(doc, "（3）運用報告書の作成及び提出")
    add_body(doc, "（4）その他前各号に付随する業務")

    # Article 3
    add_heading_text(doc, "第3条（対象資産）")
    add_body(doc, "運用対象資産: {{ managed_assets }}")

    # Article 4
    add_heading_text(doc, "第4条（報酬）")
    add_body(doc,
        "甲は、乙に対し、アセットマネジメント報酬として、運用資産残高に対し"
        "年率{{ am_fee_rate }}を乗じた金額を支払う。")
    add_body(doc,
        "報酬は四半期ごとに計算し、各四半期末日の翌月末日までに乙の指定する口座に振り込む方法により支払う。")

    # Article 5
    add_heading_text(doc, "第5条（報告義務）")
    add_body(doc,
        "乙は、甲に対し、運用状況について{{ reporting_frequency }}の頻度で報告書を提出するものとする。")
    add_body(doc,
        "甲が随時報告を求めた場合、乙は速やかにこれに応じるものとする。")

    # Article 6
    add_heading_text(doc, "第6条（善管注意義務）")
    add_body(doc,
        "乙は、善良なる管理者の注意をもって本契約に基づく業務を遂行するものとする。")

    # Article 7
    add_heading_text(doc, "第7条（期間）")
    add_body(doc,
        "本契約の有効期間は、効力発生日から1年間とする。"
        "ただし、期間満了の3ヶ月前までに甲乙いずれからも書面による終了の申出がないときは、"
        "同一条件で更に1年間更新されるものとし、以後も同様とする。")

    # Article 8
    add_heading_text(doc, "第8条（秘密保持）")
    add_body(doc,
        "甲及び乙は、本契約に関して知り得た相手方の秘密情報を、"
        "相手方の事前の書面による承諾なく第三者に開示又は漏洩してはならない。"
        "本条の義務は、本契約終了後もなお3年間存続する。")

    # Article 9
    add_heading_text(doc, "第9条（解除）")
    add_body(doc,
        "甲又は乙は、相手方が本契約の条項に違反し、相当の期間を定めて催告したにもかかわらず"
        "是正されない場合、本契約を解除することができる。")
    add_body(doc,
        "前項のほか、相手方に破産手続開始の申立てその他の重大な信用不安事由が生じた場合、"
        "催告なく直ちに本契約を解除することができる。")

    # Article 10
    add_heading_text(doc, "第10条（損害賠償）")
    add_body(doc,
        "乙は、故意又は重過失により甲に損害を与えた場合、その損害を賠償する責任を負うものとする。"
        "ただし、乙の賠償額は、直近1年間に甲が乙に支払った報酬の総額を上限とする。")

    # Article 11
    add_heading_text(doc, "第11条（準拠法及び管轄裁判所）")
    add_body(doc, "本契約は日本法に準拠し、日本法に従い解釈されるものとする。")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    # 後文
    add_blank_line(doc)
    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "asset_management_agreement.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# 8. 会計事務委託契約書（会計事務所） (Accounting Services - Firm)
# ---------------------------------------------------------------------------

def create_accounting_services_firm():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "会計事務委託契約書")

    # 前文
    add_body(doc,
        "{{ party_a_name }}（以下「甲」という）と{{ party_b_name }}（以下「乙」という）は、"
        "甲の会計事務の委託に関し、以下のとおり契約を締結する。")

    # Parties
    add_heading_text(doc, "当事者")
    add_body(doc, "甲（委託者）: {{ party_a_name }}")
    add_body(doc, "住所: {{ party_a_address }}")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（受託者・会計事務所）: {{ party_b_name }}")
    add_body(doc, "住所: {{ party_b_address }}")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")
    add_body(doc, "効力発生日: {{ effective_date }}")

    # Article 1
    add_heading_text(doc, "第1条（目的）")
    add_body(doc,
        "本契約は、甲の会計事務処理業務を乙に委託し、乙がこれを受託することを目的とする。")

    # Article 2
    add_heading_text(doc, "第2条（業務内容）")
    add_body(doc, "乙が受託する業務の範囲は以下のとおりとする。")
    add_body(doc, "業務範囲: {{ scope_of_work }}")
    add_body(doc, "（1）記帳代行及び仕訳処理")
    add_body(doc, "（2）月次試算表の作成")
    add_body(doc, "（3）決算書類の作成")
    add_body(doc, "（4）税務申告書の作成及び提出")
    add_body(doc, "（5）その他前各号に付随する業務")

    # Article 3
    add_heading_text(doc, "第3条（報酬）")
    add_body(doc,
        "甲は、乙に対し、本契約に基づく業務の対価として月額{{ monthly_fee }}円（消費税別）を支払う。")
    add_body(doc,
        "報酬は、毎月末日締めとし、翌月末日までに乙の指定する口座に振り込む方法により支払う。")

    # Article 4
    add_heading_text(doc, "第4条（資料の提供）")
    add_body(doc,
        "甲は、乙の業務遂行に必要な帳簿、証憑その他の資料を速やかに乙に提供するものとする。")
    add_body(doc,
        "甲の資料提供の遅延に起因する業務の遅延については、乙は責任を負わないものとする。")

    # Article 5
    add_heading_text(doc, "第5条（報告期限）")
    add_body(doc,
        "乙は、毎月の会計処理結果を{{ reporting_deadline }}までに甲に報告するものとする。")

    # Article 6
    add_heading_text(doc, "第6条（期間）")
    add_body(doc,
        "本契約の有効期間は、効力発生日から1年間とする。"
        "ただし、期間満了の1ヶ月前までに甲乙いずれからも書面による終了の申出がないときは、"
        "同一条件で更に1年間更新されるものとし、以後も同様とする。")

    # Article 7
    add_heading_text(doc, "第7条（秘密保持）")
    add_body(doc,
        "甲及び乙は、本契約に関して知り得た相手方の秘密情報を、"
        "相手方の事前の書面による承諾なく第三者に開示又は漏洩してはならない。"
        "本条の義務は、本契約終了後もなお3年間存続する。")

    # Article 8
    add_heading_text(doc, "第8条（解除）")
    add_body(doc,
        "甲又は乙は、相手方が本契約の条項に違反し、相当の期間を定めて催告したにもかかわらず"
        "是正されない場合、本契約を解除することができる。")

    # Article 9
    add_heading_text(doc, "第9条（免責）")
    add_body(doc,
        "乙は、甲から提供された資料の正確性について責任を負わないものとする。"
        "甲が提供した資料の誤りに起因する損害については、甲が負担するものとする。")

    # Article 10
    add_heading_text(doc, "第10条（準拠法及び管轄裁判所）")
    add_body(doc, "本契約は日本法に準拠し、日本法に従い解釈されるものとする。")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    # 後文
    add_blank_line(doc)
    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "accounting_services_firm.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# 9. 会計事務委託契約書（一般社団法人） (Accounting Services - Association)
# ---------------------------------------------------------------------------

def create_accounting_services_association():
    doc = Document()
    set_default_margins(doc)

    add_title(doc, "会計事務委託契約書")

    # 前文
    add_body(doc,
        "{{ party_a_name }}（以下「甲」という）と{{ party_b_name }}（以下「乙」という）は、"
        "甲の会計事務の委託に関し、以下のとおり契約を締結する。")

    # Parties
    add_heading_text(doc, "当事者")
    add_body(doc, "甲（委託者）: {{ party_a_name }}")
    add_body(doc, "住所: {{ party_a_address }}")
    add_body(doc, "代表者: {{ party_a_representative }}")
    add_blank_line(doc)
    add_body(doc, "乙（受託者・一般社団法人）: {{ party_b_name }}")
    add_body(doc, "住所: {{ party_b_address }}")
    add_body(doc, "代表者: {{ party_b_representative }}")
    add_blank_line(doc)
    add_body(doc, "契約日: {{ contract_date }}")
    add_body(doc, "効力発生日: {{ effective_date }}")

    # Article 1
    add_heading_text(doc, "第1条（目的）")
    add_body(doc,
        "本契約は、甲の会計事務処理業務を乙に委託し、乙がこれを受託することを目的とする。")

    # Article 2
    add_heading_text(doc, "第2条（業務内容）")
    add_body(doc, "乙が受託する業務の範囲は以下のとおりとする。")
    add_body(doc, "業務範囲: {{ scope_of_work }}")
    add_body(doc, "委託範囲: {{ delegation_scope }}")
    add_body(doc, "（1）日常の会計記帳及び仕訳処理")
    add_body(doc, "（2）月次決算の作成及び報告")
    add_body(doc, "（3）年次決算書類の作成補助")
    add_body(doc, "（4）社員総会向け会計報告資料の作成")
    add_body(doc, "（5）その他前各号に付随する業務")

    # Article 3
    add_heading_text(doc, "第3条（報酬）")
    add_body(doc,
        "甲は、乙に対し、本契約に基づく業務の対価として月額{{ monthly_fee }}円（消費税別）を支払う。")
    add_body(doc,
        "報酬は、毎月末日締めとし、翌月末日までに乙の指定する口座に振り込む方法により支払う。")

    # Article 4
    add_heading_text(doc, "第4条（資料の提供）")
    add_body(doc,
        "甲は、乙の業務遂行に必要な帳簿、証憑その他の資料を速やかに乙に提供するものとする。")
    add_body(doc,
        "甲の資料提供の遅延に起因する業務の遅延については、乙は責任を負わないものとする。")

    # Article 5
    add_heading_text(doc, "第5条（再委託）")
    add_body(doc,
        "乙は、甲の事前の書面による承諾を得た場合に限り、本契約に基づく業務の全部又は一部を"
        "第三者に再委託することができる。")
    add_body(doc,
        "再委託した場合においても、乙は本契約に基づく責任を免れないものとする。")

    # Article 6
    add_heading_text(doc, "第6条（期間）")
    add_body(doc,
        "本契約の有効期間は、効力発生日から1年間とする。"
        "ただし、期間満了の1ヶ月前までに甲乙いずれからも書面による終了の申出がないときは、"
        "同一条件で更に1年間更新されるものとし、以後も同様とする。")

    # Article 7
    add_heading_text(doc, "第7条（秘密保持）")
    add_body(doc,
        "甲及び乙は、本契約に関して知り得た相手方の秘密情報を、"
        "相手方の事前の書面による承諾なく第三者に開示又は漏洩してはならない。"
        "本条の義務は、本契約終了後もなお3年間存続する。")

    # Article 8
    add_heading_text(doc, "第8条（解除）")
    add_body(doc,
        "甲又は乙は、相手方が本契約の条項に違反し、相当の期間を定めて催告したにもかかわらず"
        "是正されない場合、本契約を解除することができる。")

    # Article 9
    add_heading_text(doc, "第9条（免責）")
    add_body(doc,
        "乙は、甲から提供された資料の正確性について責任を負わないものとする。"
        "甲が提供した資料の誤りに起因する損害については、甲が負担するものとする。")

    # Article 10
    add_heading_text(doc, "第10条（準拠法及び管轄裁判所）")
    add_body(doc, "本契約は日本法に準拠し、日本法に従い解釈されるものとする。")
    add_body(doc,
        "本契約に関する紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。")

    # 後文
    add_blank_line(doc)
    add_body(doc,
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。")

    add_signature_block(doc)

    path = os.path.join(OUTPUT_DIR, "accounting_services_association.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    create_tk_agreement()
    create_sales_agreement()
    create_master_lease()
    create_sublease_agreement()
    create_private_placement_agreement()
    create_customer_referral_agreement()
    create_asset_management_agreement()
    create_accounting_services_firm()
    create_accounting_services_association()
    print(f"\nAll 9 contract templates created in: {OUTPUT_DIR}")
